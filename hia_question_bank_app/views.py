from django.shortcuts import render, redirect
from django.core.files.storage import FileSystemStorage
from .forms import UploadFileForm
from .models import QuestionBank
from django.db.models import Max
import pandas as pd
import json
import logging
from django.http import JsonResponse
from django.db.models import Max

import os
from datetime import datetime
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image as PILImage
import io
from django.utils.html import strip_tags


from django.http import FileResponse
from django.http import HttpResponse
import os
from datetime import datetime
from docx import Document
from PIL import Image as PILImage
import io
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def upload_file(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            # Save the uploaded file temporarily
            file = request.FILES['file']
            fs = FileSystemStorage()
            filename = fs.save(file.name, file)
            uploaded_file_url = fs.url(filename)

            # Process the uploaded file
            data = pd.read_excel(fs.path(filename))

            # Replace NaN values with blank strings for text fields and 0 for numeric fields
            data = data.fillna({
                'marks': 0,
                'negative_marks': 0,
                'exam_year': 0,
            }).fillna('')

            # Get the maximum question number from the database
            max_question_number = QuestionBank.objects.aggregate(Max('question_number'))['question_number__max']
            if max_question_number:
                start_number = int(max_question_number) + 1
            else:
                start_number = 1

            # Loop through the rows and create QuestionBank entries
            for _, row in data.iterrows():
                while QuestionBank.objects.filter(question_number=str(start_number)).exists():
                    start_number += 1

                QuestionBank.objects.create(
                    question_number=str(start_number),
                    type_of_question=row.get('type_of_question', ''),
                    exam_name=row.get('exam_name', ''),
                    exam_stage=row.get('exam_stage', ''),
                    exam_year=int(row['exam_year']),
                    language=row.get('language', ''),
                    script=row.get('script', ''),
                    marks=float(row['marks']),
                    negative_marks=float(row['negative_marks']),
                    degree_of_difficulty=row.get('degree_of_difficulty', ''),
                    question_sub_type=row.get('question_sub_type', ''),
                    question_part=row.get('question_part', ''),
                    question_part_first_part=row.get('question_part_first_part', ''),
                    list_1_name=row.get('list_1_name', ''),
                    list_2_name=row.get('list_2_name', ''),
                    list_1_row1=row.get('list_1_row1', ''),
                    list_2_row1=row.get('list_2_row1', ''),
                    list_1_row2=row.get('list_1_row2', ''),
                    list_2_row2=row.get('list_2_row2', ''),
                    list_1_row3=row.get('list_1_row3', ''),
                    list_2_row3=row.get('list_2_row3', ''),
                    list_1_row4=row.get('list_1_row4', ''),
                    list_2_row4=row.get('list_2_row4', ''),
                    list_1_row5=row.get('list_1_row5', ''),
                    list_2_row5=row.get('list_2_row5', ''),
                    list_1_row6=row.get('list_1_row6', ''),
                    list_2_row6=row.get('list_2_row6', ''),
                    list_1_row7=row.get('list_1_row7', ''),
                    list_2_row7=row.get('list_2_row7', ''),
                    list_1_row8=row.get('list_1_row8', ''),
                    list_2_row8=row.get('list_2_row8', ''),
                    list_1_row9=row.get('list_1_row9', ''),
                    list_2_row9=row.get('list_2_row9', ''),
                    question_part_third_part=row.get('question_part_third_part', ''),
                    answer_option_a=row.get('answer_option_a', ''),
                    answer_option_b=row.get('answer_option_b', ''),
                    answer_option_c=row.get('answer_option_c', ''),
                    answer_option_d=row.get('answer_option_d', ''),
                    correct_answer_choice=row.get('correct_answer_choice', ''),
                    correct_answer_description=row.get('correct_answer_description', ''),
                    subject_name=row.get('subject_name', ''),
                    area_name=row.get('area_name', ''),
                    part_name=row.get('part_name', ''),
                    table_head_a=row.get('table_head_a', ''),
                    table_head_b=row.get('table_head_b', ''),
                    table_head_c=row.get('table_head_c', ''),
                    table_head_d=row.get('table_head_d', ''),
                    head_a_data1=row.get('head_a_data1', ''),
                    head_a_data2=row.get('head_a_data2', ''),
                    head_a_data3=row.get('head_a_data3', ''),
                    head_a_data4=row.get('head_a_data4', ''),
                    head_b_data1=row.get('head_b_data1', ''),
                    head_b_data2=row.get('head_b_data2', ''),
                    head_b_data3=row.get('head_b_data3', ''),
                    head_b_data4=row.get('head_b_data4', ''),
                    head_c_data1=row.get('head_c_data1', ''),
                    head_c_data2=row.get('head_c_data2', ''),
                    head_c_data3=row.get('head_c_data3', ''),
                    head_c_data4=row.get('head_c_data4', ''),
                    head_d_data1=row.get('head_d_data1', ''),
                    head_d_data2=row.get('head_d_data2', ''),
                    head_d_data3=row.get('head_d_data3', ''),
                    head_d_data4=row.get('head_d_data4', '')
                )
                start_number += 1
            return redirect('success')
    else:
        form = UploadFileForm()
    return render(request, 'hia_question_bank_app_templates/upload.html', {'form': form}) 

def success(request):
    return render(request, 'hia_question_bank_app_templates/upload_success.html')






def clean_text(text):
    """
    Clean text for writing to Word document: replace newlines with spaces and strip HTML tags.
    This function ensures that all newlines within the text are replaced with space characters
    to prevent unwanted line breaks in the Word document. HTML tags are removed to maintain 
    the purity of the text.
    """
    text = str(text)
    if text:
        # Replace all newline characters with space
        text = text.replace('\n', ' ')
        # Strip HTML tags
        cleaned_text = strip_tags(text)
        # Strip leading/trailing whitespace
        cleaned_text = cleaned_text.strip()
        return cleaned_text
    return ''

# Remove border from the List - I and List - II type Tables
def set_no_border(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['left', 'top', 'right', 'bottom']:
        border_tag = OxmlElement(f'w:{border}')
        border_tag.set(qn('w:val'), 'nil')  # Set border value to 'nil' to remove it
        border_tag.set(qn('w:sz'), '0')     # Set border size to '0' for good measure
        border_tag.set(qn('w:space'), '0')  # Set border spacing to '0'
        tcBorders.append(border_tag)
    tcPr.append(tcBorders)

# def generate_questions_document(request):
#     try:
#         # Setup directory and document file to save generated word file
#         base_dir = r'E:\Kamlesh Projects\hia_project\media\word_file'
#         os.makedirs(base_dir, exist_ok=True)
#         today = datetime.today().strftime('%Y-%m-%d')
#         file_name = f'class_plus_questions_{today}.docx'
#         file_path = os.path.join(base_dir, file_name)
#         document = Document()

#         for question in QuestionBank.objects.all():
#             # Create main table for the question and set styles
#             table = document.add_table(rows=0, cols=3)
#             table.style = 'Table Grid'

#             # Handle question text or assertion and reason
#             if question.question_part.strip():
#                 question_text = question.question_part
#             else:
#                 question_text = f"{question.question_part_first_part}\n{question.question_part_third_part}"
#                 print(question_text)


#             # Add question info to the main table
#             q_row = table.add_row().cells
#             q_row[0].text = 'Question'

#             if question.list_1_name and question.list_2_name:
#                 # Create a sub-table to handle complex question structures
#                 sub_table = document.add_table(rows=1, cols=2)
#                 sub_table.style = 'Table Grid'

#                 # Apply no border style to all cells in the sub-table as they are created
#                 for cell in sub_table._cells:
#                     set_no_border(cell)

#                 sub_hdr_cells = sub_table.rows[0].cells
#                 sub_hdr_cells[0].text = question.list_1_name
#                 sub_hdr_cells[1].text = question.list_2_name

#                 # Populate sub-table with list options
#                 for i in range(1, 10):
#                     row_cells = sub_table.add_row().cells
#                     for cell in row_cells:
#                         set_no_border(cell)  # Ensure borders are removed for new cells too
#                     list_1_option = getattr(question, f'list_1_row{i}', '')
#                     list_2_option = getattr(question, f'list_2_row{i}', '')
#                     row_cells[0].text = f"{chr(64+i)}. {list_1_option}" if list_1_option else ""
#                     row_cells[1].text = f"{i}. {list_2_option}" if list_2_option else ""

#                 # Clear the original cell content, insert question text and sub-table
#                 q_row[1]._element.clear_content()
#                 p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
#                 p.add_run(question.question_part_first_part + "\n")
#                 q_row[1]._element.append(sub_table._element)

#                 # Add 'Codes:' text below the sub-table within the same cell
#                 p = q_row[1].add_paragraph()
#                 p.add_run("\nCodes:\t A\t B\t C\t D")


#             else:
#                 # Standard question text handling
#                 q_row[1].text = question_text

#             # Merging cells for question text and image
#             q_row[1].merge(q_row[2])
            
#             # Handling image insertion if available
#             if question.image:
#                 image_path = question.image.path
#                 pil_img = PILImage.open(image_path)
#                 img_io = io.BytesIO()
#                 pil_img.save(img_io, 'JPEG')
#                 img_io.seek(0)
#                 paragraph = q_row[1].add_paragraph()
#                 run = paragraph.add_run()
#                 run.add_picture(img_io, width=Inches(1.5))

#             # Table type Questions
#             # Check for the presence of table_head_* fields and their corresponding data
#             sub_table = None

#             table_heads = ['table_head_a', 'table_head_b', 'table_head_c', 'table_head_d']
#             data_fields = [
#                 [getattr(question, f'head_a_data{j}', None) for j in range(1, 5)],
#                 [getattr(question, f'head_b_data{j}', None) for j in range(1, 5)],
#                 [getattr(question, f'head_c_data{j}', None) for j in range(1, 5)],
#                 [getattr(question, f'head_d_data{j}', None) for j in range(1, 5)]
#             ]

#             filtered_heads_data = [
#                 (head, [data for data in datas if data])
#                 for head, datas in zip(table_heads, data_fields) 
#                 if getattr(question, head, None) and any(datas)
#             ]

#             if filtered_heads_data:
#                 total_rows = 1 + max(len(datas) for _, datas in filtered_heads_data)
#                 sub_table = document.add_table(rows=total_rows, cols=len(filtered_heads_data))
#                 sub_table.style = 'Table Grid'
#                 hdr_cells = sub_table.rows[0].cells
#                 for idx, (head, _) in enumerate(filtered_heads_data):
#                     hdr_cells[idx].text = getattr(question, head, "")
#                     paragraph = hdr_cells[idx].paragraphs[0]
#                     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align header text

#                 for col_idx, (head, data_list) in enumerate(filtered_heads_data):
#                     serial_number = 1
#                     for row_idx, data in enumerate(data_list):
#                         cell = sub_table.cell(row_idx + 1, col_idx)
#                         if head == 'table_head_a':
#                             cell.text = f"{serial_number}. {data}"
#                             serial_number += 1
#                         else:
#                             cell.text = data

#             if sub_table is not None:
#                 q_row[1]._element.clear_content()
#                 p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
#                 p.add_run(question_text + "\n")
#                 q_row[1]._element.append(sub_table._element)

#             # Additional rows for type, options, solution, and marks
#             type_row = table.add_row().cells
#             type_row[0].text = 'Type'
#             type_row[1].text = question.type_of_question
#             type_row[1].merge(type_row[2])

#             correct_option_text = ""  # Store text of the correct answer
#             valid_options = ['a', 'b', 'c', 'd']  # Include all valid options

#             for opt in valid_options:
#                 option_text = getattr(question, f"answer_option_{opt}", None)  # Get the option text or None if it doesn't exist
#                 if option_text:  # Check if the option text is not None
#                     opt_row = table.add_row().cells
#                     opt_row[0].text = 'Option'
#                     opt_row[1].text = f"{opt.upper()}. {option_text}"  # Set the option text
#                     is_correct = opt == question.correct_answer_choice.lower()  # Determine if this option is correct
#                     opt_row[2].text = 'correct' if is_correct else 'incorrect'  # Set 'correct' or 'incorrect'
#                     if is_correct:
#                         correct_option_text = option_text  # Store the correct option text

#             solution_row = table.add_row().cells
#             solution_row[0].text = 'Solution'
#             solution_row[1].text = correct_option_text
#             solution_row[1].merge(solution_row[2])

#             marks_row = table.add_row().cells
#             marks_row[0].text = 'Marks'
#             marks_row[1].text = str(question.marks)
#             marks_row[2].text = str(question.negative_marks)

#             document.add_paragraph()  # Add space between questions

#         document.save(file_path)
#         return render(request, 'hia_question_bank_app_templates/generate_success.html', {'file_path': file_path})
#     except Exception as e:
#         return HttpResponse(f"An error occurred: {str(e)}", status=500)




def generate_questions_document(request):
    try:
        # Setup directory and document file to save generated word file
        base_dir = r'E:\Kamlesh Projects\hia_project\media\word_file'
        os.makedirs(base_dir, exist_ok=True)
        today = datetime.today().strftime('%Y-%m-%d')
        file_name = f'class_plus_questions_{today}.docx'
        file_path = os.path.join(base_dir, file_name)
        document = Document()

        for question in QuestionBank.objects.all():
            # Create main table for the question and set styles
            table = document.add_table(rows=0, cols=3)
            table.style = 'Table Grid'

            # Handle question text or assertion and reason
            if question.question_part and question.question_part.strip():
                question_text = question.question_part
            else:
                question_text = (question.question_part_first_part or '') + "\n" + (question.question_part_third_part or '')

            # Add question info to the main table
            q_row = table.add_row().cells
            q_row[0].text = 'Question'

            if question.list_1_name and question.list_2_name:
                # Create a sub-table to handle complex question structures
                sub_table = document.add_table(rows=1, cols=2)
                sub_table.style = 'Table Grid'

                # Apply no border style to all cells in the sub-table as they are created
                for cell in sub_table._cells:
                    set_no_border(cell)

                sub_hdr_cells = sub_table.rows[0].cells
                sub_hdr_cells[0].text = question.list_1_name
                sub_hdr_cells[1].text = question.list_2_name

                # Populate sub-table with list options
                for i in range(1, 10):
                    row_cells = sub_table.add_row().cells
                    for cell in row_cells:
                        set_no_border(cell)  # Ensure borders are removed for new cells too
                    list_1_option = getattr(question, f'list_1_row{i}', '')
                    list_2_option = getattr(question, f'list_2_row{i}', '')
                    row_cells[0].text = f"{chr(64+i)}. {list_1_option}" if list_1_option else ""
                    row_cells[1].text = f"{i}. {list_2_option}" if list_2_option else ""

                # Clear the original cell content, insert question text and sub-table
                q_row[1]._element.clear_content()
                p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
                p.add_run((question.question_part_first_part or '') + "\n")
                q_row[1]._element.append(sub_table._element)

                # Add 'Codes:' text below the sub-table within the same cell
                p = q_row[1].add_paragraph()
                p.add_run("\nCodes:\t A\t B\t C\t D")

            else:
                # Standard question text handling
                q_row[1].text = question_text

            # Merging cells for question text and image
            q_row[1].merge(q_row[2])
            
            # Handling image insertion if available
            if question.image:
                image_path = question.image.path
                pil_img = PILImage.open(image_path)
                img_io = io.BytesIO()
                pil_img.save(img_io, 'JPEG')
                img_io.seek(0)
                paragraph = q_row[1].add_paragraph()
                run = paragraph.add_run()
                run.add_picture(img_io, width=Inches(1.5))

            # Table type Questions
            # Check for the presence of table_head_* fields and their corresponding data
            sub_table = None

            table_heads = ['table_head_a', 'table_head_b', 'table_head_c', 'table_head_d']
            data_fields = [
                [getattr(question, f'head_a_data{j}', None) for j in range(1, 5)],
                [getattr(question, f'head_b_data{j}', None) for j in range(1, 5)],
                [getattr(question, f'head_c_data{j}', None) for j in range(1, 5)],
                [getattr(question, f'head_d_data{j}', None) for j in range(1, 5)]
            ]

            filtered_heads_data = [
                (head, [data for data in datas if data])
                for head, datas in zip(table_heads, data_fields) 
                if getattr(question, head, None) and any(datas)
            ]

            if filtered_heads_data:
                total_rows = 1 + max(len(datas) for _, datas in filtered_heads_data)
                sub_table = document.add_table(rows=total_rows, cols=len(filtered_heads_data))
                sub_table.style = 'Table Grid'
                hdr_cells = sub_table.rows[0].cells
                for idx, (head, _) in enumerate(filtered_heads_data):
                    hdr_cells[idx].text = getattr(question, head, "")
                    paragraph = hdr_cells[idx].paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align header text

                for col_idx, (head, data_list) in enumerate(filtered_heads_data):
                    serial_number = 1
                    for row_idx, data in enumerate(data_list):
                        cell = sub_table.cell(row_idx + 1, col_idx)
                        if head == 'table_head_a':
                            cell.text = f"{serial_number}. {data}"
                            serial_number += 1
                        else:
                            cell.text = data

            if sub_table is not None:
                q_row[1]._element.clear_content()
                p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
                p.add_run((question_text or '') + "\n")
                q_row[1]._element.append(sub_table._element)

            # Additional rows for type, options, solution, and marks
            type_row = table.add_row().cells
            type_row[0].text = 'Type'
            type_row[1].text = question.type_of_question
            type_row[1].merge(type_row[2])

            correct_option_text = ""  # Store text of the correct answer
            valid_options = ['a', 'b', 'c', 'd']  # Include all valid options
            correct_answer = question.correct_answer_choice.lower() if question.correct_answer_choice else None  # Safely handle None

            for opt in valid_options:
                option_text = getattr(question, f"answer_option_{opt}", None)  # Get the option text or None if it doesn't exist
                if option_text:  # Check if the option text is not None
                    opt_row = table.add_row().cells
                    opt_row[0].text = 'Option'
                    opt_row[1].text = f"{opt.upper()}. {option_text}"  # Set the option text
                    is_correct = opt == correct_answer  # Determine if this option is correct
                    opt_row[2].text = 'correct' if is_correct else 'incorrect'  # Set 'correct' or 'incorrect'
                    if is_correct:
                        correct_option_text = option_text  # Store the correct option text

            solution_row = table.add_row().cells
            solution_row[0].text = 'Solution'
            solution_row[1].text = correct_option_text
            solution_row[1].merge(solution_row[2])

            marks_row = table.add_row().cells
            marks_row[0].text = 'Marks'
            marks_row[1].text = str(question.marks)
            marks_row[2].text = str(question.negative_marks)

            document.add_paragraph()  # Add space between questions

        document.save(file_path)
        
        # Return the generated file as a downloadable response
        response = FileResponse(open(file_path, 'rb'), as_attachment=True, filename=file_name)
        return response

    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)


def clean_text(text):
    """
    Clean text for writing to Word document: replace newlines with spaces and strip HTML tags.
    This function ensures that all newlines within the text are replaced with space characters
    to prevent unwanted line breaks in the Word document. HTML tags are removed to maintain 
    the purity of the text.
    """
    text = str(text)
    if text:
        # Replace all newline characters with space
        text = text.replace('\n', ' ')
        # Strip HTML tags
        cleaned_text = strip_tags(text)
        # Strip leading/trailing whitespace
        cleaned_text = cleaned_text.strip()
        return cleaned_text
    return ''

def set_no_border(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['left', 'top', 'right', 'bottom']:
        border_tag = OxmlElement(f'w:{border}')
        border_tag.set(qn('w:val'), 'nil')  # Set border value to 'nil' to remove it
        border_tag.set(qn('w:sz'), '0')     # Set border size to '0' for good measure
        border_tag.set(qn('w:space'), '0')  # Set border spacing to '0'
        tcBorders.append(border_tag)
    tcPr.append(tcBorders)

# def generate_questions(request):
#     try:
#         # Setup directory and document file to save generated word file
#         base_dir = r'E:\Kamlesh Projects\hia_project\media\generate_questions'
#         os.makedirs(base_dir, exist_ok=True)
#         today = datetime.today().strftime('%Y-%m-%d')
#         file_name = f'all_questions_{today}.docx'
#         file_path = os.path.join(base_dir, file_name)
#         document = Document()

#         for question in QuestionBank.objects.all():
#             if question.question_sub_type == 'simple_type':
#                 add_simple_type(question, document)
#             elif question.question_sub_type == 'r_and_a_type':
#                 add_r_and_a_type(question, document)
#             elif question.question_sub_type == 'list_type_1':
#                 add_list_type_1(question, document)
#             elif question.question_sub_type == 'list_type_2':
#                 add_list_type_2(question, document)
            
#             # Add a space between questions
#             document.add_paragraph()

#         document.save(file_path)
#         return render(request, 'hia_question_bank_app_templates/generate_success.html', {'file_path': file_path})

#     except Exception as e:
#         return HttpResponse(f"An error occurred: {str(e)}", status=500)

# def add_simple_type(question, document):
#     """Add simple type question to the document."""
#     document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part)}")
#     add_options_and_answers(document, question)

# def add_r_and_a_type(question, document):
#     """Add reason and assertion type question to the document."""
#     document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part)}")
#     document.add_paragraph(f"{clean_text(question.question_part_third_part)}")
#     add_options_and_answers(document, question)

# def add_list_type_1(question, document):
#     """Add list type 1 question to the document."""
#     document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part_first_part)}")
#     for i in range(1, 10):
#         list_row = getattr(question, f'list_1_row{i}', None)
#         if list_row:
#             document.add_paragraph(f"{i}. {clean_text(list_row)}")
#     document.add_paragraph(f"{clean_text(question.question_part_third_part)}")
#     add_options_and_answers(document, question)

# def add_list_type_2(question, document):
#     """Add list type 2 question to the document."""
#     document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part_first_part)}")
    
#     # Create table with two columns
#     table = document.add_table(rows=1, cols=2)
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = "LIST - I"
#     hdr_cells[1].text = "LIST - II"
    
#     # Add the list names, if available
#     if question.list_1_name:
#         hdr_cells[0].text += f"\n({clean_text(question.list_1_name)})"
#     if question.list_2_name:
#         hdr_cells[1].text += f"\n({clean_text(question.list_2_name)})"

#     # Add list rows
#     labels = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
#     for i in range(1, 10):
#         list_1_item = clean_text(getattr(question, f'list_1_row{i}', ''))
#         list_2_item = clean_text(getattr(question, f'list_2_row{i}', ''))

#         if list_1_item or list_2_item:
#             row = table.add_row().cells
#             row[0].text = f"{labels[i-1]}. {list_1_item}" if list_1_item else ''
#             row[1].text = f"{i}. {list_2_item}" if list_2_item else ''

#     # Add question details and options below the table
#     if question.question_part_third_part:
#         document.add_paragraph(f"{clean_text(question.question_part_third_part)}")
    
#     add_options_and_answers(document, question)

# def add_options_and_answers(document, question):
#     """Add options and answers to the document."""
#     for opt in ['a', 'b', 'c', 'd']:
#         option_text = getattr(question, f'answer_option_{opt}', None)
#         if option_text:
#             document.add_paragraph(f"({opt.lower()}) {clean_text(option_text)}")
    
#     document.add_paragraph(f"Correct Answer: {clean_text(question.correct_answer_choice)}")
#     document.add_paragraph(f"Solution: {clean_text(question.correct_answer_description)}")




def generate_questions(request):
    try:
        # Setup directory and document file to save generated word file
        base_dir = r'E:\Kamlesh Projects\hia_project\media\generate_questions'
        os.makedirs(base_dir, exist_ok=True)
        today = datetime.today().strftime('%Y-%m-%d')
        file_name = f'all_questions_{today}.docx'
        file_path = os.path.join(base_dir, file_name)
        document = Document()

        for question in QuestionBank.objects.all():
            if question.question_sub_type == 'simple_type':
                add_simple_type(question, document)
            elif question.question_sub_type == 'r_and_a_type':
                add_r_and_a_type(question, document)
            elif question.question_sub_type == 'list_type_1':
                add_list_type_1(question, document)
            elif question.question_sub_type == 'list_type_2':
                add_list_type_2(question, document)
            
            # Add a space between questions
            document.add_paragraph()

        document.save(file_path)

        # Return the generated file as a downloadable response
        response = FileResponse(open(file_path, 'rb'), as_attachment=True, filename=file_name)
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        return response

    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)

def add_simple_type(question, document):
    """Add simple type question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part)}")
    add_options_and_answers(document, question)

def add_r_and_a_type(question, document):
    """Add reason and assertion type question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part)}")
    document.add_paragraph(f"{clean_text(question.question_part_third_part)}")
    add_options_and_answers(document, question)

def add_list_type_1(question, document):
    """Add list type 1 question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part_first_part)}")
    for i in range(1, 10):
        list_row = getattr(question, f'list_1_row{i}', None)
        if list_row:
            document.add_paragraph(f"{i}. {clean_text(list_row)}")
    document.add_paragraph(f"{clean_text(question.question_part_third_part)}")
    add_options_and_answers(document, question)

def add_list_type_2(question, document):
    """Add list type 2 question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part_first_part)}")
    
    # Create table with two columns
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "LIST - I"
    hdr_cells[1].text = "LIST - II"
    
    # Add the list names, if available
    if question.list_1_name:
        hdr_cells[0].text += f"\n({clean_text(question.list_1_name)})"
    if question.list_2_name:
        hdr_cells[1].text += f"\n({clean_text(question.list_2_name)})"

    # Add list rows
    labels = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    for i in range(1, 10):
        list_1_item = clean_text(getattr(question, f'list_1_row{i}', ''))
        list_2_item = clean_text(getattr(question, f'list_2_row{i}', ''))

        if list_1_item or list_2_item:
            row = table.add_row().cells
            row[0].text = f"{labels[i-1]}. {list_1_item}" if list_1_item else ''
            row[1].text = f"{i}. {list_2_item}" if list_2_item else ''

    # Add question details and options below the table
    if question.question_part_third_part:
        document.add_paragraph(f"{clean_text(question.question_part_third_part)}")
    
    add_options_and_answers(document, question)

def add_options_and_answers(document, question):
    """Add options and answers to the document."""
    for opt in ['a', 'b', 'c', 'd']:
        option_text = getattr(question, f'answer_option_{opt}', None)
        if option_text:
            document.add_paragraph(f"({opt.lower()}) {clean_text(option_text)}")
    
    document.add_paragraph(f"Correct Answer: {clean_text(question.correct_answer_choice)}")
    document.add_paragraph(f"Solution: {clean_text(question.correct_answer_description)}")

def clean_text(text):
    """Utility function to clean and return text, ensuring it is not None."""
    return text or ''


# Set up logging
logger = logging.getLogger(__name__)

# def add_question_view(request):
#     if request.method == 'POST':
#         try:
#             # Parsing JSON data from the request body
#             data = json.loads(request.body)

#             # Get the maximum question number from the database
#             max_question_number = QuestionBank.objects.aggregate(Max('question_number'))['question_number__max']
#             next_question_number = (int(max_question_number) + 1) if max_question_number else 1

#             # Check if provided question_no exists or if not provided, assign the next available number
#             question_no = data.get('questionNo')
#             if not question_no or QuestionBank.objects.filter(question_number=question_no).exists():
#                 while QuestionBank.objects.filter(question_number=str(next_question_number)).exists():
#                     next_question_number += 1
#                 question_no = str(next_question_number)

#             # Extracting data
#             question_type = data.get('questionType')
#             mcq_type = data.get('mcqType')
#             question_part_first = data.get('questionPartFirst')
#             reason = data.get('reason', '')
#             assertion = data.get('assertion', '')
#             question_part_second = data.get('questionPartSecond', '')
#             correct_answer = data.get('correctAnswer')
#             exam_name = data.get('examName')
#             year = data.get('year')
#             list_ii_heading1 = data.get('listIIHeading1', '')
#             list_ii_heading2 = data.get('listIIHeading2', '')
#             options = data.get('options', [])
#             list_i = data.get('listI', [])
#             list_ii = data.get('listII', [])

#             # Create and save a new QuestionBank object
#             question_bank = QuestionBank.objects.create(
#                 type_of_question=question_type,
#                 question_sub_type=mcq_type,
#                 exam_name=exam_name,
#                 exam_year=year,
#                 question_number=question_no,
#                 question_part=question_part_first,
#                 script=reason,  # Assuming 'reason' is saved under 'script'
#                 question_part_third_part=assertion,
#                 correct_answer_description=correct_answer,
#                 list_1_name=list_ii_heading1,
#                 list_2_name=list_ii_heading2,
#                 answer_option_a=options[0] if len(options) > 0 else None,
#                 answer_option_b=options[1] if len(options) > 1 else None,
#                 answer_option_c=options[2] if len(options) > 2 else None,
#                 answer_option_d=options[3] if len(options) > 3 else None,
#                 list_1_row1=list_i[0] if len(list_i) > 0 else None,
#                 list_2_row1=list_ii[0] if len(list_ii) > 0 else None,
#                 list_1_row2=list_i[1] if len(list_i) > 1 else None,
#                 list_2_row2=list_ii[1] if len(list_ii) > 1 else None,
#             )

#             return JsonResponse({'success': True, 'message': 'Question added successfully!'})
#         except Exception as e:
#             # Log the error
#             logger.error(f"Error adding question: {e}")
#             return JsonResponse({'success': False, 'message': 'Error adding question. Please try again.', 'error': str(e)})

#     return render(request, 'hia_question_bank_app_templates/add_question.html')

from django.views.decorators.csrf import csrf_exempt

# # Set up logging
# logger = logging.getLogger(__name__)

# @csrf_exempt
# def add_question_view(request):
#     if request.method == 'POST':
#         try:
#             # Parsing JSON data from the request body
#             data = json.loads(request.body)

#             # Get the maximum question number from the database
#             max_question_number = QuestionBank.objects.aggregate(Max('question_number'))['question_number__max']
#             next_question_number = (int(max_question_number) + 1) if max_question_number else 1

#             # Check if provided question_no exists or if not provided, assign the next available number
#             question_no = data.get('questionNo')
#             if not question_no or QuestionBank.objects.filter(question_number=question_no).exists():
#                 while QuestionBank.objects.filter(question_number=str(next_question_number)).exists():
#                     next_question_number += 1
#                 question_no = str(next_question_number)

#             # Extracting data
#             question_type = data.get('questionType')
#             mcq_type = data.get('mcqType')
#             question_part_first = data.get('questionPartFirst')
#             reason = data.get('reason', '')
#             assertion = data.get('assertion', '')
#             question_part_second = data.get('questionPartSecond', '')
#             correct_answer = data.get('correctAnswer')
#             answer_description = data.get('answerDescription')
#             exam_name = data.get('examName')
#             year = data.get('year')
#             marks = data.get('marks')
#             negative_marks = data.get('negativeMarks')
#             degree_of_difficulty = data.get('degreeOfDifficulty')
#             subject_name = data.get('subjectName')
#             area_name = data.get('areaName')
#             part_name = data.get('partName')
#             list_ii_heading1 = data.get('listIIHeading1', '')
#             list_ii_heading2 = data.get('listIIHeading2', '')
#             options = data.get('options', [])
#             list_i = data.get('listI', [])
#             list_ii = data.get('listII', [])

#             # Create and save a new QuestionBank object
#             question_bank = QuestionBank(
#                 type_of_question=question_type,
#                 question_sub_type=mcq_type,
#                 exam_name=exam_name,
#                 exam_year=year,
#                 question_number=question_no,
#                 reason=reason,
#                 assertion=assertion,
#                 question_part_third_part=question_part_second,
#                 correct_answer_choice=correct_answer,
#                 correct_answer_description=answer_description,
#                 marks=marks,
#                 negative_marks=negative_marks,
#                 degree_of_difficulty=degree_of_difficulty,
#                 subject_name=subject_name,
#                 area_name=area_name,
#                 part_name=part_name,
#                 list_1_name=list_ii_heading1,
#                 list_2_name=list_ii_heading2,
#                 answer_option_a=options[0] if len(options) > 0 else None,
#                 answer_option_b=options[1] if len(options) > 1 else None,
#                 answer_option_c=options[2] if len(options) > 2 else None,
#                 answer_option_d=options[3] if len(options) > 3 else None,
#                 list_1_row1=list_i[0] if len(list_i) > 0 else None,
#                 list_2_row1=list_ii[0] if len(list_ii) > 0 else None,
#                 list_1_row2=list_i[1] if len(list_i) > 1 else None,
#                 list_2_row2=list_ii[1] if len(list_ii) > 1 else None,
#                 list_1_row3=list_i[2] if len(list_i) > 2 else None,
#                 list_2_row3=list_ii[2] if len(list_ii) > 2 else None,
#                 list_1_row4=list_i[3] if len(list_i) > 3 else None,
#                 list_2_row4=list_ii[3] if len(list_ii) > 3 else None,
#                 list_1_row5=list_i[4] if len(list_i) > 4 else None,
#                 list_2_row5=list_ii[4] if len(list_ii) > 4 else None,
#                 list_1_row6=list_i[5] if len(list_i) > 5 else None,
#                 list_2_row6=list_ii[5] if len(list_ii) > 5 else None,
#                 list_1_row7=list_i[6] if len(list_i) > 6 else None,
#                 list_2_row7=list_ii[6] if len(list_ii) > 6 else None,
#                 list_1_row8=list_i[7] if len(list_i) > 7 else None,
#                 list_2_row8=list_ii[7] if len(list_ii) > 7 else None,
#                 list_1_row9=list_i[8] if len(list_i) > 8 else None,
#                 list_2_row9=list_ii[8] if len(list_ii) > 8 else None,
#             )

#             # Handle question parts based on mcqType
#             if mcq_type in ['List-I', 'List-II', 'R & A']:
#                 question_bank.question_part_first_part = question_part_first
#             else:
#                 question_bank.question_part = question_part_first

#             question_bank.save()

#             return JsonResponse({'success': True, 'message': 'Question added successfully!'})
#         except Exception as e:
#             # Log the error
#             logger.error(f"Error adding question: {e}")
#             return JsonResponse({'success': False, 'message': 'Error adding question. Please try again.', 'error': str(e)})

#     return render(request, 'hia_question_bank_app_templates/add_question.html')


# ************************ Working code ******************************************************

# # Set up logging
# logger = logging.getLogger(__name__)

# @csrf_exempt
# def add_question_view(request):
#     if request.method == 'POST':
#         try:
#             # Parsing JSON data from the request body
#             data = json.loads(request.body)

#             # Get the maximum question number from the database
#             max_question_number = QuestionBank.objects.aggregate(Max('question_number'))['question_number__max']
#             next_question_number = (int(max_question_number) + 1) if max_question_number else 1

#             # Check if provided question_no exists or if not provided, assign the next available number
#             question_no = data.get('questionNo')
#             if not question_no or QuestionBank.objects.filter(question_number=question_no).exists():
#                 while QuestionBank.objects.filter(question_number=str(next_question_number)).exists():
#                     next_question_number += 1
#                 question_no = str(next_question_number)

#             # Extracting data
#             question_type = data.get('questionType')
#             mcq_type = data.get('mcqType')
#             question_part_first = data.get('questionPartFirst')
#             reason = data.get('reason', '')
#             assertion = data.get('assertion', '')
#             question_part_second = data.get('questionPartSecond', '')
#             correct_answer = data.get('correctAnswer')
#             answer_description = data.get('answerDescription')
#             exam_name = data.get('examName')
#             year = data.get('year')
#             marks = data.get('marks')
#             negative_marks = data.get('negativeMarks')
#             degree_of_difficulty = data.get('degreeOfDifficulty')
#             subject_name = data.get('subjectName')
#             area_name = data.get('areaName')
#             part_name = data.get('partName')
#             list_ii_heading1 = data.get('listIIHeading1', '')
#             list_ii_heading2 = data.get('listIIHeading2', '')
#             options = data.get('options', [])
#             list_i = data.get('listI', [])
#             list_ii = data.get('listII', [])

#             # Create and save a new QuestionBank object
#             question_bank = QuestionBank(
#                 type_of_question=question_type,
#                 question_sub_type=mcq_type,
#                 exam_name=exam_name,
#                 exam_year=year,
#                 question_number=question_no,
#                 reason=reason,
#                 assertion=assertion,
#                 question_part_third_part=question_part_second,
#                 correct_answer_choice=correct_answer,
#                 correct_answer_description=answer_description,
#                 marks=marks,
#                 negative_marks=negative_marks,
#                 degree_of_difficulty=degree_of_difficulty,
#                 subject_name=subject_name,
#                 area_name=area_name,
#                 part_name=part_name,
#                 list_1_name=list_ii_heading1,
#                 list_2_name=list_ii_heading2,
#                 answer_option_a=options[0] if len(options) > 0 else None,
#                 answer_option_b=options[1] if len(options) > 1 else None,
#                 answer_option_c=options[2] if len(options) > 2 else None,
#                 answer_option_d=options[3] if len(options) > 3 else None,
#                 list_1_row1=list_i[0] if len(list_i) > 0 else None,
#                 list_2_row1=list_ii[0] if len(list_ii) > 0 else None,
#                 list_1_row2=list_i[1] if len(list_i) > 1 else None,
#                 list_2_row2=list_ii[1] if len(list_ii) > 1 else None,
#                 list_1_row3=list_i[2] if len(list_i) > 2 else None,
#                 list_2_row3=list_ii[2] if len(list_ii) > 2 else None,
#                 list_1_row4=list_i[3] if len(list_i) > 3 else None,
#                 list_2_row4=list_ii[3] if len(list_ii) > 3 else None,
#                 list_1_row5=list_i[4] if len(list_i) > 4 else None,
#                 list_2_row5=list_ii[4] if len(list_ii) > 4 else None,
#                 list_1_row6=list_i[5] if len(list_i) > 5 else None,
#                 list_2_row6=list_ii[5] if len(list_ii) > 5 else None,
#                 list_1_row7=list_i[6] if len(list_i) > 6 else None,
#                 list_2_row7=list_ii[6] if len(list_ii) > 6 else None,
#                 list_1_row8=list_i[7] if len(list_i) > 7 else None,
#                 list_2_row8=list_ii[7] if len(list_ii) > 7 else None,
#                 list_1_row9=list_i[8] if len(list_i) > 8 else None,
#                 list_2_row9=list_ii[8] if len(list_ii) > 8 else None,
#             )

#             # Handle question parts based on mcqType
#             if mcq_type in ['list_type_1', 'list_type_2', 'r_and_a_type']:
#                 question_bank.question_part_first_part = question_part_first
#             else:
#                 question_bank.question_part = question_part_first

#             # Save the QuestionBank object
#             question_bank.save()

#             return JsonResponse({'success': True, 'message': 'Question added successfully!'})
#         except Exception as e:
#             # Log the error
#             logger.error(f"Error adding question: {e}")
#             return JsonResponse({'success': False, 'message': 'Error adding question. Please try again.', 'error': str(e)})

#     return render(request, 'hia_question_bank_app_templates/add_question.html')














# Set up logging
logger = logging.getLogger(__name__)

@csrf_exempt
def add_question_view(request):
    if request.method == 'POST':
        try:
            # Parsing JSON data from the request body
            data = json.loads(request.body)

            # Get the maximum question number from the database
            max_question_number = QuestionBank.objects.aggregate(Max('question_number'))['question_number__max']
            next_question_number = (int(max_question_number) + 1) if max_question_number else 1

            # Check if provided question_no exists or if not provided, assign the next available number
            question_no = data.get('questionNo')
            if not question_no or QuestionBank.objects.filter(question_number=question_no).exists():
                while QuestionBank.objects.filter(question_number=str(next_question_number)).exists():
                    next_question_number += 1
                question_no = str(next_question_number)

            # Extracting data
            question_type = data.get('questionType')
            mcq_type = data.get('mcqType')
            question_part_first = data.get('questionPartFirst')
            reason = data.get('reason', '')
            assertion = data.get('assertion', '')
            question_part_second = data.get('questionPartSecond', '')
            correct_answer = data.get('correctAnswer')
            answer_description = data.get('answerDescription')
            exam_name = data.get('examName')
            year = data.get('year')
            marks = data.get('marks')
            negative_marks = data.get('negativeMarks')
            degree_of_difficulty = data.get('degreeOfDifficulty')
            subject_name = data.get('subjectName')
            area_name = data.get('areaName')
            part_name = data.get('partName')
            list_ii_heading1 = data.get('listIIHeading1', '')
            list_ii_heading2 = data.get('listIIHeading2', '')
            options = data.get('options', [])
            list_i = data.get('listI', [])
            list_ii = data.get('listII', [])

            # Create and save a new QuestionBank object
            question_bank = QuestionBank(
                type_of_question=question_type,
                question_sub_type=mcq_type,
                exam_name=exam_name,
                exam_year=year,
                question_number=question_no,
                reason=reason,
                assertion=assertion,
                question_part_third_part=question_part_second,
                correct_answer_choice=correct_answer,
                correct_answer_description=answer_description,
                marks=marks,
                negative_marks=negative_marks,
                degree_of_difficulty=degree_of_difficulty,
                subject_name=subject_name,
                area_name=area_name,
                part_name=part_name,
                list_1_name=list_ii_heading1,
                list_2_name=list_ii_heading2,
                answer_option_a=options[0] if len(options) > 0 else None,
                answer_option_b=options[1] if len(options) > 1 else None,
                answer_option_c=options[2] if len(options) > 2 else None,
                answer_option_d=options[3] if len(options) > 3 else None,
                list_1_row1=list_i[0] if len(list_i) > 0 else None,
                list_2_row1=list_ii[0] if len(list_ii) > 0 else None,
                list_1_row2=list_i[1] if len(list_i) > 1 else None,
                list_2_row2=list_ii[1] if len(list_ii) > 1 else None,
                list_1_row3=list_i[2] if len(list_i) > 2 else None,
                list_2_row3=list_ii[2] if len(list_ii) > 2 else None,
                list_1_row4=list_i[3] if len(list_i) > 3 else None,
                list_2_row4=list_ii[3] if len(list_ii) > 3 else None,
                list_1_row5=list_i[4] if len(list_i) > 4 else None,
                list_2_row5=list_ii[4] if len(list_ii) > 4 else None,
                list_1_row6=list_i[5] if len(list_i) > 5 else None,
                list_2_row6=list_ii[5] if len(list_ii) > 5 else None,
                list_1_row7=list_i[6] if len(list_i) > 6 else None,
                list_2_row7=list_ii[6] if len(list_ii) > 6 else None,
                list_1_row8=list_i[7] if len(list_i) > 7 else None,
                list_2_row8=list_ii[7] if len(list_ii) > 7 else None,
                list_1_row9=list_i[8] if len(list_i) > 8 else None,
                list_2_row9=list_ii[8] if len(list_ii) > 8 else None,
            )

            # Handle question parts based on mcqType
            if mcq_type in ['List-I', 'List-II', 'R & N']:
                question_bank.question_part_first_part = question_part_first
            else:
                question_bank.question_part = question_part_first

            # Save the QuestionBank object
            question_bank.save()

            return JsonResponse({'success': True, 'message': 'Question added successfully!'})
        except Exception as e:
            # Log the error
            logger.error(f"Error adding question: {e}")
            return JsonResponse({'success': False, 'message': 'Error adding question. Please try again.', 'error': str(e)})

    return render(request, 'hia_question_bank_app_templates/add_question.html')